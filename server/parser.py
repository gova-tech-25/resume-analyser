import re
from pypdf import PdfReader
from docx import Document
import io

class ParsingError(Exception):
    def __init__(self, code, message):
        self.code = code
        self.message = message
        super().__init__(message)

def sanitize_text(text: str) -> str:
    # 1. Strip script tags
    text = re.sub(r'<script\b[^<]*(?:(?!<\/script>)<[^<]*)*<\/script>', '', text, flags=re.IGNORECASE)
    
    # 2. Mask Emails
    text = re.sub(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', '[email@masked.com]', text)
    
    # 3. Mask Phones
    # Matches various formats: +1-123-456-7890, (123) 456-7890, 123.456.7890, 1234567890
    text = re.sub(r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', '[phone masked]', text)
    
    # 4. Mask SSNs
    text = re.sub(r'\b\d{3}-\d{2}-\d{4}\b', '[SSN masked]', text)
    
    return text

def parse_pdf(file_bytes: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        
        # Check if password protected
        if reader.is_encrypted:
            raise ParsingError("password_protected", "This PDF is password-protected. Please remove the password and re-upload.")
            
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
                
        text = text.strip()
        if not text:
            # Check if pdf has pages but no text (scanned PDF)
            if len(reader.pages) > 0:
                raise ParsingError("scanned_pdf", "This PDF appears to be a scan. Please upload a text-based PDF or copy-paste your resume text.")
            else:
                raise ParsingError("file_empty", "The file appears to be empty.")
                
        return text
    except ParsingError:
        raise
    except Exception as e:
        # Check if PDF header is missing or file is corrupted
        err_msg = str(e).lower()
        if "eof marker not found" in err_msg or "not a valid pdf" in err_msg or "pdf header" in err_msg:
            raise ParsingError("file_parse", "The PDF file is corrupted or invalid. Try saving it again and re-uploading.")
        raise ParsingError("file_parse", f"Failed to parse PDF: {str(e)}")

def parse_docx(file_bytes: bytes) -> str:
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text.append(paragraph.text)
                
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text.append(cell.text)
                        
        extracted_text = "\n".join(text).strip()
        if not extracted_text:
            raise ParsingError("file_empty", "The file appears to be empty.")
            
        return extracted_text
    except ParsingError:
        raise
    except Exception as e:
        raise ParsingError("corrupted_docx", f"We couldn't read this file. Try saving as .docx again or paste your resume as text.")

def parse_txt(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode('utf-8', errors='ignore').strip()
        if not text:
            raise ParsingError("file_empty", "The file appears to be empty.")
        return text
    except ParsingError:
        raise
    except Exception as e:
        raise ParsingError("file_parse", f"Failed to parse TXT file: {str(e)}")

def process_resume_text(text: str) -> dict:
    # Basic word count
    words = text.split()
    word_count = len(words)
    
    if word_count == 0:
        raise ParsingError("file_empty", "The file appears to be empty.")
        
    is_short_resume = word_count < 200
    is_non_resume = word_count < 100
    
    # Simple language check heuristic:
    # If the ratio of non-ASCII alphabetic characters is > 15%, flag it
    alpha_chars = [c for c in text if c.isalpha()]
    non_ascii_alpha = [c for c in alpha_chars if ord(c) > 127]
    is_non_english = len(non_ascii_alpha) / len(alpha_chars) > 0.15 if alpha_chars else False
    
    # Truncate to first 3000 words if very long
    is_truncated = False
    if word_count > 3000:
        text = " ".join(words[:3000])
        is_truncated = True
        word_count = 3000
        
    # Sanitize and mask PII
    sanitized_text = sanitize_text(text)
    
    return {
        "text": sanitized_text,
        "word_count": word_count,
        "is_short_resume": is_short_resume,
        "is_non_resume": is_non_resume,
        "is_non_english": is_non_english,
        "is_truncated": is_truncated
    }
